import os
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq

load_dotenv()
my_api_key = os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("api key not found")

client = Groq(api_key=my_api_key)

model = "llama-3.3-70b-versatile"

job_des = """
Description
Do you want to solve real customer problems through innovative technology? Do you enjoy working on scalable services in a collaborative team environment? Do you want to see your code directly impact millions of customers worldwide?

At Amazon, we hire the best minds in technology to innovate and build on behalf of our customers. Customer obsession is part of our company DNA, which has made us one of the world's most beloved brands.

Our Software Development Engineers (SDEs) use modern technology to solve complex problems while seeing their work's impact first-hand. The challenges SDEs solve at Amazon are meaningful and influence millions of customers, sellers, and products globally. We seek individuals passionate about creating new products, features, and services while managing ambiguity in an environment where development cycles are measured in weeks, not years.

At Amazon, we believe in ownership at every level. As an SDE-I, you'll own the entire lifecycle of your code - from design through deployment and ongoing operations. This ownership mindset, combined with our commitment to operational excellence, ensures we deliver the highest quality solutions for our customers.

We're looking for curious minds who think big and want to define tomorrow's technology. At Amazon, you'll grow into the high-impact engineer you know you can be, supported by a culture of learning and mentorship. Every day brings exciting new challenges and opportunities for personal growth.
Key job responsibilities
• Collaborate and communicate effectively with experienced cross-disciplinary Amazonians to design, build, and operate innovative products and services that delight our customers, while participating in technical discussions to drive solutions forward.
• Design and develop scalable solutions using cloud-native architectures and microservices in a large distributed computing environment.
• Participate in code reviews and contribute to technical documentation.
• Build and maintain resilient distributed systems that are scalable, fault-tolerant, and cost-effective.
• Leverage and contribute to the development of GenAI and AI-powered tools to enhance development productivity while staying current with emerging technologies.
• Write clean, maintainable code following best practices and design patterns.
• Work in an agile environment practicing CI/CD principles while participating in operational responsibilities including on-call duties.
• Demonstrate operational excellence through monitoring, troubleshooting, and resolving production issues.
Basic Qualifications
- Experience with at least one general-purpose programming language such as Java, Python, C++, C#, Go, Rust, or TypeScript
- Experience with data structure implementation, basic algorithm development, and/or object-oriented design principles
- Currently has, or is in the process of obtaining a bachelor’s degree in Computer Science, Computer Engineering, Data Science, Information Systems, or related STEM fields
- Must be 18 years of age of older
Preferred Qualifications
- Experience from previous technical internship(s) or demonstrated project experience
- Experience with one or more of the following: AI tools for development productivity, Cloud platforms (preferably AWS), Database systems (SQL and NoSQL), Contributing to open-source projects, Version control systems, Debugging and troubleshooting complex systems
- Demonstrated ability to learn and adapt to new technologies quickly
- Basic understanding of software development lifecycle (SDLC)
- Strong problem-solving and analytical skills
- Excellent written and verbal communication skills
"""
from pydantic import BaseModel, Field
class JobD(BaseModel):
    role : str
    required_skills : list[str]
    preferred_skills: list[str]
    minimum_experience : float | None
    education_requirements : list[str]
    responsibilities : list[str]

jobd_schema = JobD.model_json_schema()

system_prompt = f"""
you are an expert HR assistant.

your job is to analyze job description and extract structured information from them.

return only valid json matching this schema:

{jobd_schema}
IMPORTANT:
do not return the schema itself.
do not return fields like "properties", "titles" or "type".
fill the schema with actual information extracted from the job description.

if minimum experience is not mentioned, return null.
if information for a list is missing, return an empty list.
do not invent information
"""

user_prompt = f"""
analyze the following job description:

{job_des}
"""
message_system= {
    "role" : "system",
    "content" : system_prompt
}
message_user ={
    "role" : "user",
    "content" : user_prompt
}
response_format={
    "type" : "json_object"
}


messages = [message_system, message_user]

response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)

answer = response.choices[0].message.content

raw_json = answer

import json

job_data = json.loads(raw_json)

job = JobD(**job_data)

print(job.minimum_experience)
print(job.education_requirements)

class MatchResult(BaseModel):
    score : float
    details : dict

class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: str | None = None

class Resume(BaseModel):
    name : str | None = None
    email : str | None = None
    phone : str | None = None

    total_experience_years : float | None = None

    skills: list[str] = []
    experience: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certification: list[str] = []

resume_schema = Resume.model_json_schema()

def final_score(job, resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    you are a HR recruiter.

    compare the candidate's resume with the job description.

    job description:
    {job.model_dump_json(indent=2)}

    candidate resume:
    {resume.model_dump_json(indent=2)}

    return json matching this schema:
    {match_schema}

    give me:
    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """

    message = {
        "role": "user",
        "content": prompt
    }

    messages =[message]
    response_format={
        "type": "json_object"
    }

    response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)

    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)

def parse_resume(resume_text):
    system_prompt = f"""
    you are an expert resume parser.

    extract information from the resume based on its meaning, not only based on exact section headings.

    different resumes may use different headings.

    for example: 
    - experience
    - professional experience
    - work history
    - employment
    - internships

    these may all contain relevant experience.

    skills may also appear in the skills section, work experience, internships or projects.

    return only valid json matching this schema:

    {resume_schema}

    important rules:
    
    1. do not invent information.
    2. if a value is not available, return null.
    3. if a list has no informationl, return as empty list.
    4. include interships inside experiences.
    5. extract skills mentioned across the entire resume.

    """

    user_prompt = f"""
    parse the following resume:

    {resume_text}
    """

    message_system={
        "role" : "system",
        "content" : system_prompt
    }

    message_user={
        "role" : "user",
        "content" : user_prompt
    }

    messages=[message_system, message_user]
    response_format={
        "type": "json_object"
    }

    response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)

    raw_output = response.choices[0].message.content

    data = json.loads(raw_output)

    resume = Resume(**data)
    return resume

from pypdf import PdfReader
from docx import Document

def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"

    return text


def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"

    return text

def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)

    else:
        return None

import time
resume_folder = Path("resumes")
all_results = []
for file_path in resume_folder.iterdir():
    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue

    print("\nProcessing:", file_path.name)
    resume_text = read_resume(file_path)
    parsed_resume = parse_resume(resume_text)
    time.sleep(5)
    result = final_score(job, parsed_resume)

    time.sleep(5)

    print("score: ", result.score)
    all_results.append({
        "name": parsed_resume.name,
        "score": result.score,
        "details": result.details
    })

all_results.sort(
    key=lambda candidate: candidate["score"],
    reverse=True
)

top2 = all_results[:2]
worst2 = all_results[-2:]

print("top 2 candidates")
for candidate in top2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )

    print(candidate["details"])

print("lowest 2 candidates")
for candidate in worst2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )

    print(candidate["details"])


